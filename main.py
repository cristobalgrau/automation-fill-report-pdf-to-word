import pypdf
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# docx Documentation: https://python-docx.readthedocs.io/en/latest/index.html
# pypdf Documentation: https://pypdf.readthedocs.io/en/stable/


# ========================== FUNCTIONS TO WORK WITH PDF DOCUMENT ==========================

def extract_text_from_pdf(pdf_path):
    with open(pdf_path, "rb") as pdf:
        reader = pypdf.PdfReader(pdf)
        page = reader.pages[0].extract_text()
    return page.splitlines()


def clean_text_list(text_list):
    text_list = [line for line in text_list if line.strip()]  # Remove empty lines
    return text_list[4:12]  # Get relevant lines


def extract_values(new_list):
    # Extract and clean values
    attendance_date = new_list[0].replace("<XXXXXXXXXXXXXXXX>  Attendance Date", "").strip()
    location = new_list[1].replace("Attendance Location", "").strip()
    ticket = new_list[2].replace("Customer", "").strip()
    antenna_serial = new_list[6].replace("Antenna Serial", "").strip()
    antenna_model = new_list[7].replace("Antenna type", "").strip()

    # Obtain WO and Case
    temp = new_list[3].replace("Click or tap here to enter text.", "").replace("Case / WO", "").lstrip()
    wo_complete = temp.split()[0]
    case = temp.replace(wo_complete, "").strip()
    case = case.strip(" ()")
    wo = wo_complete.replace("WO-", "000")

    # Obtain Vessel Name
    temp = new_list[4].replace("Request Date", "")
    request_date = temp.split()[0]
    vessel = temp.replace(request_date, "").lstrip().replace("Vessel Name", "").strip()

    placeholders = {
        "attendance_date": attendance_date,
        "[location]": location,
        "[ticket]": ticket,
        "[wo]": wo,
        "[case]": case,
        "[vessel]": vessel,
        "[antenna_serial]": antenna_serial,
        "[antenna_model]": antenna_model,
    }

    file_name = f"{wo_complete}_Field Service Report_{vessel}_date.docx"

    return placeholders, file_name


def print_extracted_values(values):
    print("DATA FROM PDF:")
    for key, value in values.items():
        print(f"{key.replace('_', ' ').strip("[]").title()}: {value}")
    print("\nDATA REPLACED IN REPORT:")


# ========================== FUNCTIONS TO WORK WITH WORD DOCUMENT ==========================

def replace_text_and_align(cell, placeholder, replacement):
    if cell.text == placeholder:
        cell.text = cell.text.replace(placeholder, replacement)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        print(f"Found {placeholder} and modified!")


def replace_paragraph_text(doc, placeholder, replacement):
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, replacement)
            print(f"Found {placeholder} and modified!")


# ========================== MAIN ==========================

# Working with the PDF
pdf_path = "file_name.pdf"
text_list = extract_text_from_pdf(pdf_path)
new_list = clean_text_list(text_list)

replacements, file_name = extract_values(new_list)
print_extracted_values(replacements)


# Working with the DOCX
doc_path = 'report.docx'
output_path = file_name

# Load the document
doc = Document(doc_path)

# Set the default font and size
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(10)

# replace placeholder in paragraph for attendance date
attendance_date_key = "attendance_date"
replace_paragraph_text(doc, attendance_date_key, replacements[attendance_date_key])

# Iterate through tables and replace placeholders
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for placeholder, replacement in replacements.items():
                replace_text_and_align(cell, placeholder, replacement)

# Save the modified document
doc.save(output_path)

